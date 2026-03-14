"""
MILEAN Telegram Bot — Юридический ИИ-ассистент
Full-featured: AI chat, file analysis, RAG, web search, instruction slots
"""

import asyncio
import json
import logging
import os
import re
import ssl
import tempfile
import time
from io import BytesIO
from typing import Optional

import aiohttp
from aiogram import Bot, Dispatcher, F, Router, types
from aiogram import BaseMiddleware
from aiogram.enums import ParseMode, ChatAction
from aiogram.filters import Command, CommandStart
from aiogram.types import (
    BotCommand, FSInputFile, InlineKeyboardButton,
    InlineKeyboardMarkup, Message, CallbackQuery
)
from aiogram.fsm.context import FSMContext
from aiogram.fsm.state import State, StatesGroup
from aiogram.fsm.storage.memory import MemoryStorage

try:
    import sentry_sdk  # type: ignore
except Exception:
    sentry_sdk = None

try:
    import redis  # type: ignore
except Exception:
    redis = None

# ─── CONFIG ───
BOT_TOKEN = os.getenv("BOT_TOKEN", "").strip()
BOT_TOKENS_ENV = os.getenv("BOT_TOKENS", "").strip()
NVIDIA_URL = "https://integrate.api.nvidia.com/v1/chat/completions"
NVIDIA_KEY = os.getenv("NVIDIA_API_KEY", "").strip()
PPLX_KEY = os.getenv("PPLX_API_KEY", "").strip()
PPLX_URL = "https://api.perplexity.ai/chat/completions"
MODEL_HEAVY = "qwen/qwen3.5-397b-a17b"   # For documents + Think
MODEL = "meta/llama-3.1-405b-instruct"     # Primary fast
MODEL_FAST = "meta/llama-3.3-70b-instruct" # Fallback fast
PPLX_MODEL = "sonar"                       # Perplexity with internet
MAX_HISTORY = 10
MAX_FILE_SIZE = 20 * 1024 * 1024  # 20MB
SYNC_API = "https://milean.vercel.app/api/sync"
CODEX_API = "https://milean.vercel.app/api/codex"
CODEX_DEFAULT_TOKEN = os.getenv("CODEX_DEFAULT_TOKEN", "").strip().upper()
REDIS_URL = os.getenv("REDIS_URL", "").strip()
USER_STATE_TTL_SEC = int(os.getenv("USER_STATE_TTL_SEC", "2592000"))  # 30 days
SENTRY_DSN = os.getenv("SENTRY_DSN", "").strip()
ALERT_BOT_TOKEN = os.getenv("ALERT_BOT_TOKEN", "").strip()
ALERT_CHAT_ID = os.getenv("ALERT_CHAT_ID", "").strip()
CHUNK_SIZE = 500
RAG_TOP_K = 8

logging.basicConfig(level=logging.INFO)
log = logging.getLogger("milean-bot")
_redis_client = None
_alert_last_sent = {}

if sentry_sdk and SENTRY_DSN:
    try:
        sentry_sdk.init(dsn=SENTRY_DSN, traces_sample_rate=0.1)
    except Exception as e:
        log.warning(f"Sentry init failed: {e}")


LEGAL_SOURCE_DOMAINS = [
    "consultant.ru",
    "garant.ru",
    "publication.pravo.gov.ru",
    "pravo.gov.ru",
    "duma.gov.ru",
    "ksrf.ru",
    "vsrf.ru",
]


def _get_redis():
    global _redis_client
    if not REDIS_URL or redis is None:
        return None
    if _redis_client is not None:
        return _redis_client
    try:
        _redis_client = redis.Redis.from_url(
            REDIS_URL,
            decode_responses=True,
            socket_connect_timeout=2,
            socket_timeout=2,
            retry_on_timeout=True,
        )
        _redis_client.ping()
    except Exception as e:
        log.warning(f"Redis unavailable: {e}")
        _redis_client = None
    return _redis_client


def _user_state_key(uid: int, bot_id: Optional[int] = None) -> str:
    if bot_id is None:
        return f"milean:tguser:unknown:{uid}"
    return f"milean:tguser:{bot_id}:{uid}"


def _serialize_user(u: dict) -> dict:
    return {
        "hist": u.get("hist", []),
        "files": u.get("files", []),
        "chunks": u.get("chunks", []),
        "instr": u.get("instr", ""),
        "web_on": bool(u.get("web_on", False)),
        "think_on": bool(u.get("think_on", False)),
        "active_slot": u.get("active_slot", "milean"),
        "slots": u.get("slots", {}),
        "project_name": u.get("project_name", ""),
        "project_token": u.get("project_token", ""),
        "last_response": u.get("last_response", ""),
        "last_thinking": u.get("last_thinking", ""),
        "last_query": u.get("last_query", ""),
        "last_codex_task": u.get("last_codex_task", ""),
    }


def _load_user_from_redis(uid: int, bot_id: Optional[int]) -> Optional[dict]:
    c = _get_redis()
    if not c:
        return None
    try:
        raw = c.get(_user_state_key(uid, bot_id))
        if not raw:
            return None
        d = json.loads(raw)
        return d if isinstance(d, dict) else None
    except Exception:
        return None


def _save_user_to_redis(uid: int, bot_id: Optional[int], u: dict) -> None:
    c = _get_redis()
    if not c:
        return
    try:
        c.setex(
            _user_state_key(uid, bot_id),
            USER_STATE_TTL_SEC,
            json.dumps(_serialize_user(u), ensure_ascii=False),
        )
    except Exception:
        pass


async def _send_alert(where: str, message: str, min_interval_sec: int = 120) -> None:
    if not ALERT_BOT_TOKEN or not ALERT_CHAT_ID:
        return
    key = f"{where}:{message[:140]}"
    now = int(time.time())
    prev = _alert_last_sent.get(key, 0)
    if now - prev < min_interval_sec:
        return
    _alert_last_sent[key] = now
    text = f"⚠️ MILEAN bot error [{where}]\n{message[:3000]}"
    url = f"https://api.telegram.org/bot{ALERT_BOT_TOKEN}/sendMessage"
    payload = {"chat_id": ALERT_CHAT_ID, "text": text}
    timeout = aiohttp.ClientTimeout(total=8, connect=3)
    try:
        async with aiohttp.ClientSession(timeout=timeout) as session:
            async with session.post(url, data=payload):
                pass
    except Exception:
        pass


def _report_exception(where: str, exc: Exception, extra: Optional[dict] = None) -> None:
    log.error(f"{where}: {exc}", exc_info=True)
    if sentry_sdk:
        try:
            with sentry_sdk.push_scope() as scope:
                scope.set_tag("component", where)
                if extra:
                    for k, v in extra.items():
                        scope.set_extra(str(k), v)
                sentry_sdk.capture_exception(exc)
        except Exception:
            pass
    try:
        loop = asyncio.get_running_loop()
        loop.create_task(_send_alert(where, str(exc)))
    except RuntimeError:
        pass


def _collect_bot_tokens() -> list[str]:
    """Collect unique bot tokens from BOT_TOKENS (comma-separated) and BOT_TOKEN."""
    tokens = []
    if BOT_TOKENS_ENV.strip():
        tokens.extend(t.strip() for t in BOT_TOKENS_ENV.split(",") if t.strip())
    if BOT_TOKEN.strip():
        tokens.append(BOT_TOKEN.strip())
    unique = []
    seen = set()
    for token in tokens:
        if token not in seen:
            seen.add(token)
            unique.append(token)
    return unique


BOT_TOKENS = _collect_bot_tokens()


def _validate_env() -> None:
    missing = []
    if not BOT_TOKENS:
        missing.append("BOT_TOKEN/BOT_TOKENS")
    if not NVIDIA_KEY:
        missing.append("NVIDIA_API_KEY")
    if missing:
        raise RuntimeError("Missing required env vars: " + ", ".join(missing))
    if not PPLX_KEY:
        log.warning("PPLX_API_KEY is not set; web search will use DuckDuckGo fallback only.")
    if not REDIS_URL:
        log.warning("REDIS_URL is not set; user state will be in-memory only.")
    if not SENTRY_DSN:
        log.warning("SENTRY_DSN is not set; exception monitoring is limited to logs.")

# ─── DEFAULT INSTRUCTION ───
MILEAN_INSTR = """Ты — высококлассный российский юрист и адвокат с практикой более 20 лет.
Специализация:
— мошенничество с недвижимостью
— подделка подписей и документов
— злоупотребление доверенностями
— отчуждение долей, ПДКП, ничтожные сделки
— гражданские и уголовные дела (в т.ч. 159, 160, 327 УК РФ)

🎯 ЦЕЛЬ РАБОТЫ
Комплексно разобрать все эпизоды дела.
Выявить и зафиксировать:
• ничтожность и оспоримость сделок
• подделку подписей и использование факсимиле без полномочий
• злоупотребление доверенностями
• незаконное распоряжение денежными средствами

⚖️ ОБЯЗАТЕЛЬНЫЕ ИСТОЧНИКИ ПРАВА
ГК РФ, УК РФ (159, 160, 165, 327), ФЗ-218, Постановления Пленума ВС РФ

🧠 ПРАВИЛА РАБОТЫ
• Жёстко разделяй: факт / правовая оценка / гипотеза
• Любой вывод — через норму права
• Не допускать фантазий и домыслов

📊 ФОРМАТ ОТВЕТА
1️⃣ КРАТКИЙ ВЫВОД (5–10 строк)
2️⃣ ТАБЛИЦА НАРУШЕНИЙ
3️⃣ УГОЛОВНО-ПРАВОВАЯ ОЦЕНКА
4️⃣ ЭКСПЕРТИЗЫ
5️⃣ ПРОЦЕССУАЛЬНЫЕ ШАГИ"""

# ─── USER DATA STORAGE (in-memory) ───
users = {}  # f"{bot_id}:{user_id}" -> {hist, files, chunks, instr, ...}


def get_user(uid: int, bot_id: Optional[int] = None) -> dict:
    key = f"{bot_id}:{uid}" if bot_id is not None else str(uid)
    if key not in users:
        base = {
            "hist": [],
            "files": [],  # [{name, chunks, chars}]
            "chunks": [],  # [{text, file}]
            "instr": MILEAN_INSTR,
            "web_on": False,
            "think_on": False,
            "active_slot": "milean",
            "slots": {},  # {slot_id: {name, text}}
            "project_name": "",
            "project_token": "",
            "last_response": "",
            "last_thinking": "",
            "last_query": "",
            "last_codex_task": "",
        }
        restored = _load_user_from_redis(uid, bot_id)
        if isinstance(restored, dict):
            base.update(restored)
        users[key] = base
    return users[key]


def get_project_token(u: dict) -> str:
    """Prefer /connect token from chat state; fallback to service-wide token."""
    token = (u.get("project_token") or "").strip().upper()
    if token:
        return token
    return CODEX_DEFAULT_TOKEN


# ─── TEXT EXTRACTION ───
def extract_pdf(data: bytes) -> str:
    from PyPDF2 import PdfReader
    reader = PdfReader(BytesIO(data))
    text = ""
    for page in reader.pages:
        t = page.extract_text()
        if t:
            text += t + "\n"
    return text


def extract_docx(data: bytes) -> str:
    from docx import Document
    doc = Document(BytesIO(data))
    return "\n".join(p.text for p in doc.paragraphs if p.text.strip())


def extract_txt(data: bytes) -> str:
    for enc in ("utf-8", "cp1251", "latin-1"):
        try:
            return data.decode(enc)
        except:
            continue
    return data.decode("utf-8", errors="replace")


def extract_text(filename: str, data: bytes) -> str:
    ext = filename.lower().rsplit(".", 1)[-1] if "." in filename else ""
    if ext == "pdf":
        return extract_pdf(data)
    elif ext in ("docx", "doc"):
        return extract_docx(data)
    else:
        return extract_txt(data)


# ─── CHUNKING & RAG ───
def chunk_text(text: str, size: int = CHUNK_SIZE) -> list:
    chunks = []
    sentences = re.split(r'(?<=[.!?。\n])\s+', text)
    cur = ""
    for s in sentences:
        if len(cur + s) > size and cur:
            chunks.append(cur.strip())
            cur = ""
        cur += s + " "
    if cur.strip():
        chunks.append(cur.strip())
    return chunks


def local_search(query: str, chunks: list, top_k: int = RAG_TOP_K) -> list:
    if not chunks:
        return []
    q_words = re.sub(r'[^\w\s]', '', query.lower()).split()
    q_words = [w for w in q_words if len(w) > 2]
    if not q_words:
        return chunks[:top_k]

    scored = []
    for ch in chunks:
        txt = ch["text"].lower()
        score = 0
        for w in q_words:
            idx = 0
            while True:
                idx = txt.find(w, idx)
                if idx == -1:
                    break
                score += 1
                idx += len(w)
        if score > 0:
            scored.append((ch, score))

    scored.sort(key=lambda x: -x[1])
    return [s[0] for s in scored[:top_k]]


# ─── WEB SEARCH KEYWORDS ───
_WEB_KEYWORDS = [
    "погода", "weather", "новости", "news", "курс", "цена", "сколько стоит",
    "сегодня", "завтра", "вчера", "сейчас", "актуальн", "последн", "свежи",
    "2024", "2025", "2026", "расписание", "результат матча", "счёт",
    "где купить", "где найти", "как доехать", "адрес", "телефон", "сайт",
    "что случилось", "что произошло", "кто выиграл", "кто победил",
    "рецепт", "отзыв", "рейтинг", "топ ", "лучший", "обзор",
]


def _needs_web(query: str) -> bool:
    """Auto-detect if query needs internet search."""
    q = query.lower()
    for kw in _WEB_KEYWORDS:
        if kw in q:
            return True
    # Question words + no files context → likely needs web
    if any(q.startswith(w) for w in ["как ", "что ", "где ", "когда ", "почему ", "какой ", "какая ", "какое ", "кто "]):
        return True
    return False


# ─── PERPLEXITY SEARCH (with internet) ───
def _is_legal_query(query: str) -> bool:
    q = (query or "").lower()
    return bool(re.search(r"\b(гк|ук|гпк|апк|нк|коап|фз|ст\.?)\b|закон|поправ|изменени|редакц|пленум|постановлен|правов|юрид|судеб|consultant|garant|права", q))


def _extract_domains(text: str) -> list[str]:
    out = []
    seen = set()
    if not text:
        return out
    for m in re.finditer(r"((?:https?://)?(?:www\.)?[a-z0-9][a-z0-9.-]+\.[a-z]{2,})(?:/[^\s]*)?", text.lower()):
        d = (m.group(1) or "").strip().replace("https://", "").replace("http://", "")
        d = d.split("/")[0]
        if d.startswith("www."):
            d = d[4:]
        if d and "." in d and d not in seen:
            seen.add(d)
            out.append(d)
    return out


def _legal_domains_from_instruction(instr: str) -> list[str]:
    out = []
    seen = set()

    def add(d: str):
        if d and d not in seen:
            seen.add(d)
            out.append(d)

    for d in _extract_domains(instr or ""):
        add(d)
    for d in LEGAL_SOURCE_DOMAINS:
        add(d)
    return out[:6]


async def perplexity_search(query: str, domains: Optional[list[str]] = None) -> tuple[str, list[str]]:
    """Use Perplexity API for internet-powered answers."""
    if not PPLX_KEY:
        return "", []
    try:
        ssl_ctx = ssl.create_default_context()
        ssl_ctx.check_hostname = False
        ssl_ctx.verify_mode = ssl.CERT_NONE

        payload = {
            "model": PPLX_MODEL,
            "messages": [
                {"role": "system", "content": "Отвечай кратко и по делу на русском языке. Давай актуальную информацию с источниками."},
                {"role": "user", "content": query}
            ],
            "max_tokens": 1024,
            "stream": False,
            "return_citations": True,
        }
        if domains:
            payload["search_domain_filter"] = domains[:6]
        headers = {
            "Content-Type": "application/json",
            "Authorization": f"Bearer {PPLX_KEY}"
        }

        timeout = aiohttp.ClientTimeout(total=30, connect=10)
        async with aiohttp.ClientSession() as session:
            async with session.post(
                PPLX_URL, json=payload, headers=headers,
                ssl=ssl_ctx, timeout=timeout
            ) as resp:
                if resp.status != 200:
                    log.warning(f"Perplexity {resp.status}")
                    return "", []
                data = await resp.json()
                content = data["choices"][0]["message"].get("content", "")
                citations = data.get("citations", []) or []
                if content:
                    return f"🌐 Информация из интернета:\n{content}", citations
                return "", citations
    except Exception as e:
        log.error(f"Perplexity error: {e}")
        return "", []


# ─── DUCKDUCKGO SEARCH (fallback) ───
async def web_search(query: str) -> tuple[str, list[str]]:
    try:
        from duckduckgo_search import DDGS
        results = []
        citations = []
        with DDGS() as ddgs:
            for r in ddgs.text(query, max_results=5):
                results.append(f"• {r['title']}: {r['body']}")
                href = (r.get("href") or r.get("url") or "").strip()
                if href and href not in citations:
                    citations.append(href)
        if results:
            return "🌐 Результаты веб-поиска:\n" + "\n".join(results), citations
        return "", citations
    except Exception as e:
        log.error(f"Web search error: {e}")
        return "", []


async def smart_web_search(query: str, instr_text: str = "") -> tuple[str, list[str], str]:
    """Try Perplexity first, fallback to DuckDuckGo.
    Returns (context, citations, checked_at_utc).
    """
    checked_at = time.strftime("%Y-%m-%d %H:%M:%S UTC", time.gmtime())
    legal_mode = _is_legal_query(query)
    domains = _legal_domains_from_instruction(instr_text) if legal_mode else []

    # Force legal query refinement for freshest legal amendments.
    q = query
    if legal_mode and not re.search(r"поправ|изменени|редакц|на сегодня|действующ", q.lower()):
        q = f"{query} актуальная редакция на сегодня"

    context_parts = []
    citations = []
    seen = set()

    def add_citations(arr: list[str]):
        for u in arr or []:
            u = (u or "").strip()
            if not u or u in seen:
                continue
            seen.add(u)
            citations.append(u)

    # 1) Perplexity general / legal-domain-filtered
    pplx_ctx, pplx_citations = await perplexity_search(q, domains=domains if domains else None)
    if pplx_ctx:
        context_parts.append(pplx_ctx)
        add_citations(pplx_citations)

    # 2) DDG fallback
    if not context_parts:
        ddg_ctx, ddg_citations = await web_search(q)
        if ddg_ctx:
            context_parts.append(ddg_ctx)
            add_citations(ddg_citations)

    # 3) For legal mode, run targeted domain queries for amendments.
    if legal_mode and domains:
        for d in domains[:3]:
            ddg_ctx, ddg_citations = await web_search(f"{q} site:{d}")
            if ddg_ctx:
                context_parts.append(f"⚖️ Юр-поиск по {d}:\n{ddg_ctx}")
                add_citations(ddg_citations)

    if not context_parts:
        return "", [], checked_at

    lines = ["\n\n".join(context_parts), f"🗓 Дата проверки источников: {checked_at}"]
    if citations:
        lines.append("🔗 Источники:\n" + "\n".join(f"- {u}" for u in citations[:12]))
    return "\n\n".join(lines), citations, checked_at


async def codex_enqueue_task(token: str, task: str, msg: Message) -> dict:
    payload = {
        "token": token,
        "action": "enqueue",
        "task": task,
        "chat_id": msg.chat.id,
        "user_id": msg.from_user.id,
        "username": msg.from_user.username or "",
        "source": "telegram",
    }
    ssl_ctx = ssl.create_default_context()
    ssl_ctx.check_hostname = False
    ssl_ctx.verify_mode = ssl.CERT_NONE
    timeout = aiohttp.ClientTimeout(total=25, connect=10)
    async with aiohttp.ClientSession() as session:
        async with session.post(
            CODEX_API, json=payload, ssl=ssl_ctx, timeout=timeout
        ) as resp:
            data = await resp.json()
            if resp.status != 200 or not data.get("ok"):
                raise RuntimeError(data.get("error") or f"HTTP {resp.status}")
            return data.get("task") or {}


async def codex_list_tasks(token: str, limit: int = 10) -> list:
    ssl_ctx = ssl.create_default_context()
    ssl_ctx.check_hostname = False
    ssl_ctx.verify_mode = ssl.CERT_NONE
    timeout = aiohttp.ClientTimeout(total=25, connect=10)
    url = f"{CODEX_API}?token={token}&action=list&limit={max(1,min(limit,50))}"
    async with aiohttp.ClientSession() as session:
        async with session.get(url, ssl=ssl_ctx, timeout=timeout) as resp:
            data = await resp.json()
            if resp.status != 200 or not data.get("ok"):
                raise RuntimeError(data.get("error") or f"HTTP {resp.status}")
            return data.get("tasks") or []


# ─── NVIDIA API CALL ───
async def _call_model(model: str, messages: list, max_tokens: int = 4096,
                      timeout_sec: int = 60) -> tuple:
    """Call a single model, return (content, thinking) or raise."""
    ssl_ctx = ssl.create_default_context()
    ssl_ctx.check_hostname = False
    ssl_ctx.verify_mode = ssl.CERT_NONE

    payload = {
        "model": model,
        "messages": messages,
        "temperature": 0.6,
        "top_p": 0.95,
        "max_tokens": max_tokens,
        "stream": True
    }

    headers = {
        "Content-Type": "application/json",
        "Authorization": f"Bearer {NVIDIA_KEY}"
    }

    timeout = aiohttp.ClientTimeout(total=timeout_sec, connect=10)
    async with aiohttp.ClientSession() as session:
        async with session.post(
            NVIDIA_URL, json=payload, headers=headers,
            ssl=ssl_ctx, timeout=timeout
        ) as resp:
            if resp.status != 200:
                err = await resp.text()
                raise Exception(f"API {resp.status}: {err[:200]}")

            content = ""
            thinking = ""
            async for line in resp.content:
                line = line.decode("utf-8").strip()
                if not line or not line.startswith("data: "):
                    continue
                chunk_str = line[6:]
                if chunk_str == "[DONE]":
                    break
                try:
                    chunk = json.loads(chunk_str)
                    delta = chunk["choices"][0].get("delta", {})
                    if "content" in delta and delta["content"]:
                        content += delta["content"]
                    if "reasoning_content" in delta and delta["reasoning_content"]:
                        thinking += delta["reasoning_content"]
                except (json.JSONDecodeError, KeyError, IndexError):
                    continue

            return content, thinking


async def call_nvidia(messages: list, think: bool = False, has_docs: bool = False) -> tuple:
    """Smart model selection:
    - Documents + Think → heavy qwen3.5-397b (deep analysis)
    - Normal → fast 405b → fallback 70b
    """
    if think and has_docs:
        # Heavy model for document analysis with thinking
        log.info("Using HEAVY model (qwen3.5-397b) for document analysis")
        try:
            return await _call_model(MODEL_HEAVY, messages, max_tokens=8192, timeout_sec=180)
        except Exception as e:
            log.warning(f"Heavy model failed: {e}, falling back to primary...")

    # Fast path: primary → fallback
    try:
        return await _call_model(MODEL, messages, max_tokens=4096, timeout_sec=60)
    except Exception as e:
        log.warning(f"Primary model failed: {e}, trying fast model...")

    try:
        return await _call_model(MODEL_FAST, messages, max_tokens=4096, timeout_sec=45)
    except Exception as e:
        raise Exception(f"Серверы недоступны. Попробуйте позже.") from e


# ─── TELEGRAM BOT ───
storage = MemoryStorage()
dp = Dispatcher(storage=storage)
router = Router()


class PersistUserMiddleware(BaseMiddleware):
    async def __call__(self, handler, event, data):
        result = await handler(event, data)
        try:
            from_user = getattr(event, "from_user", None)
            bot_obj = data.get("bot") or getattr(event, "bot", None)
            if from_user and bot_obj:
                u = get_user(from_user.id, bot_obj.id)
                _save_user_to_redis(from_user.id, bot_obj.id, u)
        except Exception:
            pass
        return result


dp.update.middleware(PersistUserMiddleware())


class EditInstr(StatesGroup):
    waiting = State()


class SaveSlot(StatesGroup):
    waiting_name = State()
    slot_id = State()


# ─── COMMANDS ───
@router.message(CommandStart())
async def cmd_start(msg: Message):
    u = get_user(msg.from_user.id, msg.bot.id)
    kb = get_main_keyboard(u)
    await msg.answer(
        "⚖️ <b>MILEAN — Юридическая Компания</b>\n\n"
        "Я ваш юридический ИИ-ассистент.\n\n"
        "📎 Отправьте файлы (PDF, DOCX, TXT) для анализа\n"
        "💬 Задайте вопрос — я отвечу с учётом загруженных документов\n"
        "🌐 Включите веб-поиск для актуальной информации\n\n"
        "Используйте /help для списка команд\n"
        "Используйте /panel для панели управления",
        parse_mode=ParseMode.HTML,
        reply_markup=kb
    )


@router.message(Command("help"))
async def cmd_help(msg: Message):
    await msg.answer(
        "📋 <b>Команды MILEAN:</b>\n\n"
        "⚖️ /milean — загрузить стандартную инструкцию\n"
        "📝 /instr — показать текущую инструкцию\n"
        "✏️ /setinstr — написать свою инструкцию\n"
        "💾 /saveslot <i>N</i> — сохранить инструкцию в слот (1-10)\n"
        "📌 /loadslot <i>N</i> — загрузить инструкцию из слота\n"
        "📋 /slots — список слотов\n"
        "🗑 /clearinstr — очистить инструкцию\n\n"
        "📎 /files — список загруженных файлов\n"
        "🗑 /clearfiles — удалить все файлы\n\n"
        "🔗 /connect <i>TOKEN</i> — подключить проект с web\n"
        "🔑 /token — как получить токен\n\n"
        "🤖 /codex <i>задача</i> — отправить задачу в Codex relay\n"
        "📡 /codexstatus — статус очереди Codex\n\n"
        "🌐 /web — вкл/выкл веб-поиск\n"
        "🧠 /think — вкл/выкл режим Think\n"
        "🔄 /clear — очистить историю чата\n"
        "⚙️ /settings — текущие настройки\n"
        "🎛 /panel — панель управления с кнопками\n"
        "❓ /help — эта справка",
        parse_mode=ParseMode.HTML
    )


@router.message(Command("panel"))
async def cmd_panel(msg: Message):
    u = get_user(msg.from_user.id, msg.bot.id)
    await msg.answer(_get_status_text(u), parse_mode=ParseMode.HTML, reply_markup=get_main_keyboard(u))


@router.message(Command("milean"))
async def cmd_milean(msg: Message):
    u = get_user(msg.from_user.id, msg.bot.id)
    u["instr"] = MILEAN_INSTR
    u["active_slot"] = "milean"
    await msg.answer("⚖️ Инструкция MILEAN загружена!", reply_markup=get_main_keyboard(u))


@router.message(Command("instr"))
async def cmd_instr(msg: Message):
    u = get_user(msg.from_user.id, msg.bot.id)
    instr = u["instr"]
    if instr:
        text = f"📝 <b>Текущая инструкция</b> ({u['active_slot']}):\n\n{instr[:3000]}"
        if len(instr) > 3000:
            text += "\n\n<i>...обрезано</i>"
    else:
        text = "📝 Инструкция не задана. Используйте /milean или /setinstr"
    await msg.answer(text, parse_mode=ParseMode.HTML)


@router.message(Command("setinstr"))
async def cmd_setinstr(msg: Message, state: FSMContext):
    await state.set_state(EditInstr.waiting)
    await msg.answer(
        "✏️ Отправьте текст новой инструкции.\n"
        "Или /cancel для отмены.",
        parse_mode=ParseMode.HTML
    )


@router.message(EditInstr.waiting)
async def process_setinstr(msg: Message, state: FSMContext):
    if msg.text and msg.text.startswith("/cancel"):
        await state.clear()
        await msg.answer("❌ Отменено")
        return
    u = get_user(msg.from_user.id, msg.bot.id)
    u["instr"] = msg.text or ""
    u["active_slot"] = "custom"
    await state.clear()
    await msg.answer("✅ Инструкция обновлена!", reply_markup=get_main_keyboard(u))


@router.message(Command("clearinstr"))
async def cmd_clearinstr(msg: Message):
    u = get_user(msg.from_user.id, msg.bot.id)
    u["instr"] = ""
    u["active_slot"] = "empty"
    await msg.answer("🗑 Инструкция очищена", reply_markup=get_main_keyboard(u))


@router.message(Command("saveslot"))
async def cmd_saveslot(msg: Message):
    u = get_user(msg.from_user.id, msg.bot.id)
    parts = (msg.text or "").split(maxsplit=2)
    if len(parts) < 2:
        await msg.answer("Использование: /saveslot <i>N</i> <i>название</i>\nПример: /saveslot 1 Бухгалтерия", parse_mode=ParseMode.HTML)
        return
    try:
        slot_id = int(parts[1])
        if not 1 <= slot_id <= 10:
            raise ValueError
    except:
        await msg.answer("❌ Номер слота от 1 до 10")
        return
    if not u["instr"]:
        await msg.answer("❌ Инструкция пуста — нечего сохранять")
        return
    name = parts[2] if len(parts) > 2 else f"Слот {slot_id}"
    u["slots"][str(slot_id)] = {"name": name, "text": u["instr"]}
    await msg.answer(f"💾 Сохранено в слот {slot_id}: <b>{name}</b>", parse_mode=ParseMode.HTML)


@router.message(Command("loadslot"))
async def cmd_loadslot(msg: Message):
    u = get_user(msg.from_user.id, msg.bot.id)
    parts = (msg.text or "").split()
    if len(parts) < 2:
        await msg.answer("Использование: /loadslot <i>N</i>", parse_mode=ParseMode.HTML)
        return
    try:
        slot_id = int(parts[1])
    except:
        await msg.answer("❌ Укажите номер слота (1-10)")
        return
    slot = u["slots"].get(str(slot_id))
    if not slot:
        await msg.answer(f"❌ Слот {slot_id} пуст")
        return
    u["instr"] = slot["text"]
    u["active_slot"] = f"slot_{slot_id}"
    await msg.answer(f"📌 Загружен слот {slot_id}: <b>{slot['name']}</b>", parse_mode=ParseMode.HTML, reply_markup=get_main_keyboard(u))


@router.message(Command("slots"))
async def cmd_slots(msg: Message):
    u = get_user(msg.from_user.id, msg.bot.id)
    lines = ["📋 <b>Слоты инструкций:</b>\n"]
    for i in range(1, 11):
        slot = u["slots"].get(str(i))
        if slot:
            lines.append(f"  {i}. 📌 <b>{slot['name']}</b> ({len(slot['text'])} симв.)")
        else:
            lines.append(f"  {i}. <i>пусто</i>")
    lines.append(f"\n⚖️ MILEAN: {'✅ активна' if u['active_slot'] == 'milean' else 'доступна'}")
    await msg.answer("\n".join(lines), parse_mode=ParseMode.HTML)


@router.message(Command("web"))
async def cmd_web(msg: Message):
    u = get_user(msg.from_user.id, msg.bot.id)
    u["web_on"] = not u["web_on"]
    status = "✅ включён" if u["web_on"] else "❌ выключен"
    await msg.answer(f"🌐 Веб-поиск: {status}", reply_markup=get_main_keyboard(u))


@router.message(Command("think"))
async def cmd_think(msg: Message):
    u = get_user(msg.from_user.id, msg.bot.id)
    u["think_on"] = not u["think_on"]
    status = "✅ включён" if u["think_on"] else "❌ выключен"
    await msg.answer(f"🧠 Think режим: {status}", reply_markup=get_main_keyboard(u))


@router.message(Command("clear"))
async def cmd_clear(msg: Message):
    u = get_user(msg.from_user.id, msg.bot.id)
    u["hist"] = []
    await msg.answer("🔄 История чата очищена", reply_markup=get_main_keyboard(u))


@router.message(Command("files"))
async def cmd_files(msg: Message):
    u = get_user(msg.from_user.id, msg.bot.id)
    if not u["files"]:
        await msg.answer("📎 Нет загруженных файлов.\nОтправьте PDF, DOCX или TXT для анализа.")
        return
    lines = ["📎 <b>Загруженные файлы:</b>\n"]
    total_chunks = 0
    total_chars = 0
    for f in u["files"]:
        lines.append(f"  📄 <b>{f['name']}</b> — {f['chunks']} чанков, {f['chars']} симв.")
        total_chunks += f["chunks"]
        total_chars += f["chars"]
    lines.append(f"\n📊 Всего: {len(u['files'])} файлов, {total_chunks} чанков, {total_chars} символов")
    await msg.answer("\n".join(lines), parse_mode=ParseMode.HTML)


@router.message(Command("clearfiles"))
async def cmd_clearfiles(msg: Message):
    u = get_user(msg.from_user.id, msg.bot.id)
    u["files"] = []
    u["chunks"] = []
    await msg.answer("🗑 Все файлы удалены", reply_markup=get_main_keyboard(u))


@router.message(Command("settings"))
async def cmd_settings(msg: Message):
    u = get_user(msg.from_user.id, msg.bot.id)
    slot_name = u["active_slot"]
    if slot_name == "milean":
        slot_name = "⚖️ MILEAN"
    elif slot_name == "empty":
        slot_name = "⛔ Пусто"
    elif slot_name.startswith("slot_"):
        sid = slot_name.split("_")[1]
        s = u["slots"].get(sid, {})
        slot_name = f"📌 {s.get('name', 'Слот ' + sid)}"
    else:
        slot_name = "✏️ Своя"

    await msg.answer(
        "⚙️ <b>Настройки:</b>\n\n"
        f"📝 Инструкция: {slot_name}\n"
        f"🧠 Think: {'✅' if u['think_on'] else '❌'}\n"
        f"🌐 Веб-поиск: {'✅' if u['web_on'] else '❌'}\n"
        f"📎 Файлов: {len(u['files'])}\n"
        f"💬 История: {len(u['hist'])//2} сообщений\n"
        f"📊 Чанков в RAG: {len(u['chunks'])}",
        parse_mode=ParseMode.HTML,
        reply_markup=get_main_keyboard(u)
    )


@router.message(Command("cancel"))
async def cmd_cancel(msg: Message, state: FSMContext):
    await state.clear()
    await msg.answer("❌ Отменено")


@router.message(Command("connect"))
async def cmd_connect(msg: Message):
    u = get_user(msg.from_user.id, msg.bot.id)
    parts = (msg.text or "").split()
    if len(parts) < 2:
        await msg.answer(
            "🔗 <b>Подключение проекта с Web</b>\n\n"
            "Использование: <code>/connect TOKEN</code>\n\n"
            "Как получить токен:\n"
            "1. Откройте milean.vercel.app\n"
            "2. В разделе «Проекты» нажмите 📤 рядом с проектом\n"
            "3. Скопируйте токен и вставьте сюда",
            parse_mode=ParseMode.HTML
        )
        return

    token = parts[1].strip().upper()
    status_msg = await msg.answer(f"🔗 Подключение проекта <code>{token}</code>...", parse_mode=ParseMode.HTML)

    try:
        ssl_ctx = ssl.create_default_context()
        ssl_ctx.check_hostname = False
        ssl_ctx.verify_mode = ssl.CERT_NONE

        async with aiohttp.ClientSession() as session:
            async with session.get(
                f"{SYNC_API}?token={token}",
                ssl=ssl_ctx,
                timeout=aiohttp.ClientTimeout(total=30)
            ) as resp:
                data = await resp.json()

        if not data.get("ok"):
            await status_msg.edit_text(
                f"❌ Проект не найден.\n\n"
                f"Убедитесь что:\n"
                f"• Токен верный: <code>{token}</code>\n"
                f"• Вы нажали 📤 Sync на web-сайте\n"
                f"• Прошло менее 24 часов с синхронизации",
                parse_mode=ParseMode.HTML
            )
            return

        # Load project data
        u["instr"] = data.get("instr", "")
        u["active_slot"] = "web_project"
        u["chunks"] = data.get("chunks", [])
        u["files"] = data.get("files", [])
        u["hist"] = data.get("hist", [])
        u["project_name"] = data.get("name", "Web Project")
        u["project_token"] = token

        proj_name = data.get("name", "Web Project")
        files_count = len(u["files"])
        chunks_count = len(u["chunks"])
        hist_count = len(u["hist"]) // 2

        await status_msg.edit_text(
            f"✅ <b>Проект подключён!</b>\n\n"
            f"📂 <b>{proj_name}</b>\n"
            f"📎 Файлов: {files_count}\n"
            f"🧩 Чанков: {chunks_count}\n"
            f"💬 Сообщений: {hist_count}\n"
            f"📝 Инструкция: {'✅ загружена' if u['instr'] else '❌ нет'}\n\n"
            f"🔑 Токен: <code>{token}</code>\n\n"
            f"Теперь просто отправьте вопрос — я отвечу с учётом "
            f"файлов и инструкций из вашего web-проекта!",
            parse_mode=ParseMode.HTML,
            reply_markup=get_main_keyboard(u)
        )

    except Exception as e:
        _report_exception("bot.connect", e, {"token": token[:16] if 'token' in locals() else ""})
        await status_msg.edit_text(f"❌ Ошибка подключения: {e}")


@router.message(Command("token"))
async def cmd_token(msg: Message):
    """Show info about how to get token"""
    await msg.answer(
        "🔑 <b>Токен проекта</b>\n\n"
        "Каждый проект на milean.vercel.app имеет уникальный токен.\n"
        "Формат может отличаться (например <code>ML-XXXXXXXX</code> или UUID).\n\n"
        "📤 Чтобы перенести проект в Telegram:\n"
        "1. Откройте <b>Проекты</b> на сайте\n"
        "2. Нажмите 📤 рядом с нужным проектом\n"
        "3. Скопируйте токен\n"
        "4. Отправьте <code>/connect TOKEN</code>\n\n"
        "⚡️ Синхронизируются: инструкция, файлы, история чата",
        parse_mode=ParseMode.HTML
    )


@router.message(Command("codex"))
async def cmd_codex(msg: Message):
    """Send task to Codex relay queue for current connected project."""
    u = get_user(msg.from_user.id, msg.bot.id)
    token = get_project_token(u)
    if not token:
        await msg.answer(
            "❌ <b>Сначала подключите проект:</b>\n"
            "<code>/connect TOKEN</code>\n\n"
            "После подключения команда /codex будет отправлять задачи "
            "в общую очередь Codex relay.",
            parse_mode=ParseMode.HTML
        )
        return

    parts = (msg.text or "").split(maxsplit=1)
    if len(parts) < 2 or not parts[1].strip():
        await cmd_codexstatus(msg)
        return

    task = parts[1].strip()
    if len(task) > 6000:
        task = task[:6000]
    u["last_codex_task"] = task
    wait_msg = await msg.answer("⏳ Отправка задачи в очередь Codex...")
    try:
        t = await codex_enqueue_task(token, task, msg)
        tid = t.get("id", "—")
        await wait_msg.edit_text(
            "✅ <b>Задача поставлена в очередь Codex</b>\n\n"
            f"🆔 <code>{tid}</code>\n"
            f"📂 Проект: <code>{token}</code>\n"
            f"📝 Задача:\n<blockquote>{_escape(task[:900] + ('...' if len(task) > 900 else ''))}</blockquote>\n\n"
            "Проверьте статус: /codexstatus",
            parse_mode=ParseMode.HTML
        )
    except Exception as e:
        await wait_msg.edit_text(f"❌ Ошибка очереди Codex: {e}")


@router.message(Command("codexstatus"))
async def cmd_codexstatus(msg: Message):
    """Show Codex relay queue status for connected project."""
    u = get_user(msg.from_user.id, msg.bot.id)
    token = get_project_token(u)
    if not token:
        await msg.answer(
            "❌ Проект не подключён.\n"
            "Сначала: <code>/connect TOKEN</code>",
            parse_mode=ParseMode.HTML
        )
        return
    try:
        tasks = await codex_list_tasks(token, limit=8)
    except Exception as e:
        await msg.answer(f"❌ Не удалось получить статус Codex: {e}")
        return
    if not tasks:
        await msg.answer(
            f"📡 <b>Очередь Codex ({token}) пуста</b>\n"
            "Отправьте задачу: <code>/codex исправь ...</code>",
            parse_mode=ParseMode.HTML
        )
        return

    icons = {
        "queued": "🕒",
        "claimed": "🛠",
        "done": "✅",
        "error": "❌",
        "skipped": "⏭",
        "canceled": "🚫",
    }
    lines = [f"📡 <b>Codex очередь проекта {token}</b>\n"]
    for t in tasks:
        st = (t.get("status") or "queued").lower()
        icon = icons.get(st, "•")
        tid = t.get("id", "—")
        txt = (t.get("task") or "").strip().replace("\n", " ")
        if len(txt) > 90:
            txt = txt[:90] + "..."
        who = t.get("claimed_by") or t.get("completed_by") or ""
        if who:
            lines.append(f"{icon} <code>{tid}</code> · <b>{st}</b> · {who}\n<blockquote>{_escape(txt)}</blockquote>")
        else:
            lines.append(f"{icon} <code>{tid}</code> · <b>{st}</b>\n<blockquote>{_escape(txt)}</blockquote>")
    await msg.answer("\n".join(lines), parse_mode=ParseMode.HTML)


# ─── FILE HANDLER ───
@router.message(F.document)
async def handle_document(msg: Message):
    u = get_user(msg.from_user.id, msg.bot.id)
    doc = msg.document

    if doc.file_size > MAX_FILE_SIZE:
        await msg.answer(f"❌ Файл слишком большой (макс. {MAX_FILE_SIZE // 1024 // 1024}MB)")
        return

    ext = (doc.file_name or "").lower().rsplit(".", 1)[-1] if doc.file_name and "." in doc.file_name else ""
    supported = ("pdf", "docx", "doc", "txt", "py", "js", "json", "csv", "md", "html", "xml", "log")
    if ext not in supported:
        await msg.answer(f"❌ Формат .{ext} не поддерживается.\nПоддерживаются: {', '.join(supported)}")
        return

    status_msg = await msg.answer(f"📄 <b>{doc.file_name}</b>\n⏳ Загрузка и анализ...", parse_mode=ParseMode.HTML)
    await msg.bot.send_chat_action(msg.chat.id, ChatAction.TYPING)

    try:
        file = await msg.bot.get_file(doc.file_id)
        data = BytesIO()
        await msg.bot.download_file(file.file_path, data)
        file_bytes = data.getvalue()

        text = extract_text(doc.file_name, file_bytes)
        if not text.strip():
            await status_msg.edit_text(f"📄 <b>{doc.file_name}</b>\n⚠️ Не удалось извлечь текст", parse_mode=ParseMode.HTML)
            return

        chunks = chunk_text(text)
        for ch in chunks:
            u["chunks"].append({"text": ch, "file": doc.file_name})

        fobj = {"name": doc.file_name, "chunks": len(chunks), "chars": len(text)}
        u["files"].append(fobj)

        await status_msg.edit_text(
            f"✅ <b>{doc.file_name}</b> обработан!\n"
            f"📊 {len(chunks)} чанков · {len(text)} символов\n"
            f"📎 Всего файлов: {len(u['files'])}",
            parse_mode=ParseMode.HTML
        )

    except Exception as e:
        _report_exception("bot.file_processing", e, {"file": doc.file_name if 'doc' in locals() else ""})
        await status_msg.edit_text(f"❌ Ошибка обработки: {e}", parse_mode=ParseMode.HTML)


# ─── MAIN MESSAGE HANDLER ───
@router.message(F.text & ~F.text.startswith("/"))
async def handle_message(msg: Message, state: FSMContext):
    u = get_user(msg.from_user.id, msg.bot.id)
    query = msg.text.strip()
    if not query:
        return

    await msg.bot.send_chat_action(msg.chat.id, ChatAction.TYPING)
    status_msg = await msg.answer("⏳ Этап 1/4: подготовка запроса...")

    try:
        # Build system message
        sys_parts = []
        web_citations = []
        web_checked_at = ""
        legal_mode = _is_legal_query(query)

        # Instruction
        if u["instr"]:
            sys_parts.append(u["instr"])

        # Web search — auto or manual
        use_web = u["web_on"] or _needs_web(query)
        has_docs = bool(u["chunks"])

        if use_web:
            await status_msg.edit_text("🌐 Этап 2/4: поиск и проверка источников...")
            web_context, web_citations, web_checked_at = await smart_web_search(query, u.get("instr", ""))
            if web_context:
                sys_parts.append(web_context)
                sys_parts.append(
                    "При формировании ответа используй факты только из найденных веб-источников выше. "
                    "Укажи дату проверки и не делай категоричных выводов без подтверждающих ссылок."
                )
            elif legal_mode:
                sys_parts.append(
                    "Не удалось получить подтвержденные свежие источники по правовому запросу. "
                    "Прямо сообщи об этом и перечисли, какие источники нужно проверить вручную."
                )

        # RAG search
        if u["chunks"]:
            await status_msg.edit_text("🔍 Этап 2/4: RAG поиск по файлам...")
            relevant = local_search(query, u["chunks"])
            if relevant:
                rag_parts = []
                for ch in relevant:
                    rag_parts.append(f"[{ch['file']}]: {ch['text']}")
                rag_context = "📎 Контекст из файлов:\n" + "\n---\n".join(rag_parts)
                sys_parts.append(rag_context)

        # Status message based on mode
        if u["think_on"] and has_docs:
            await status_msg.edit_text("🧠 Этап 3/4: глубокий анализ документов...")
        else:
            await status_msg.edit_text("🧠 Этап 3/4: генерация ответа...")

        # Build messages
        messages = []
        if sys_parts:
            messages.append({"role": "system", "content": "\n\n".join(sys_parts)})

        # Add history
        for h in u["hist"][-MAX_HISTORY:]:
            messages.append(h)

        messages.append({"role": "user", "content": query})

        # Call AI — smart model selection
        content, thinking = await call_nvidia(messages, think=u["think_on"], has_docs=has_docs)

        # Save to history
        if web_citations:
            src_lines = "\n".join(f"• {u}" for u in web_citations[:10])
            stamp = web_checked_at or time.strftime("%Y-%m-%d %H:%M:%S UTC", time.gmtime())
            content = (
                content.rstrip()
                + "\n\n📎 Источники проверки:\n"
                + src_lines
                + f"\n🗓 Дата проверки: {stamp}"
            )
        elif legal_mode:
            content = (
                content.rstrip()
                + "\n\n⚠️ Примечание: не удалось получить подтвержденные свежие юридические источники в этом сеансе. "
                  "Проверьте consultant.ru / garant.ru / pravo.gov.ru вручную."
            )

        u["hist"].append({"role": "user", "content": query})
        u["hist"].append({"role": "assistant", "content": content})

        # Trim history
        if len(u["hist"]) > MAX_HISTORY * 2:
            u["hist"] = u["hist"][-MAX_HISTORY * 2:]

        # Save to user for file generation
        await status_msg.edit_text("✅ Этап 4/4: ответ готов и отправлен")

        # Short preview in chat + full response as file
        preview = content[:300].replace("<", "&lt;").replace(">", "&gt;")
        if len(content) > 300:
            preview += "..."

        # Thinking block
        think_text = ""
        if thinking and u["think_on"]:
            think_short = thinking[:300] + "..." if len(thinking) > 300 else thinking
            think_text = f"💭 <b>Размышление:</b>\n<blockquote>{_escape(think_short)}</blockquote>\n\n"

        # Send preview + file buttons
        preview_msg = think_text + f"📄 <b>Ответ:</b>\n{preview}\n\n⬇️ Скачайте полный ответ:"
        file_kb = InlineKeyboardMarkup(inline_keyboard=[
            [
                InlineKeyboardButton(text="📄 TXT", callback_data=f"dl_txt_{msg.message_id}"),
                InlineKeyboardButton(text="📝 DOCX", callback_data=f"dl_docx_{msg.message_id}"),
                InlineKeyboardButton(text="📕 PDF", callback_data=f"dl_pdf_{msg.message_id}"),
            ],
            [
                InlineKeyboardButton(text="💬 Показать в чате", callback_data=f"dl_chat_{msg.message_id}"),
            ]
        ])

        # Store response for download
        u["last_response"] = content
        u["last_thinking"] = thinking or ""
        u["last_query"] = query

        try:
            await msg.bot.send_message(
                msg.chat.id, preview_msg,
                parse_mode=ParseMode.HTML,
                reply_to_message_id=msg.message_id,
                reply_markup=file_kb
            )
        except:
            await msg.bot.send_message(
                msg.chat.id,
                f"📄 Ответ готов ({len(content)} симв.)\n⬇️ Скачайте полный ответ:",
                reply_to_message_id=msg.message_id,
                reply_markup=file_kb
            )

    except asyncio.TimeoutError as e:
        _report_exception("bot.handle_message.timeout", e, {"query": query[:200]})
        try:
            await status_msg.edit_text("❌ Таймаут — NVIDIA API не ответил. Попробуйте ещё раз или отключите Think (/think)")
        except:
            pass
    except Exception as e:
        err_msg = str(e) or type(e).__name__
        _report_exception("bot.handle_message.error", e, {"query": query[:200]})
        try:
            await status_msg.edit_text(f"❌ Ошибка: {err_msg[:500]}")
        except:
            pass


# ─── HELPER: update keyboard safely ───
async def _update_kb(cb: CallbackQuery, u: dict, toast: str):
    """Update keyboard on message, ignore 'not modified' errors"""
    kb = get_main_keyboard(u)
    status = _get_status_text(u)
    await cb.answer(toast)
    try:
        await cb.message.edit_text(status, reply_markup=kb, parse_mode=ParseMode.HTML)
    except Exception as e:
        if "not modified" in str(e).lower():
            return  # Same content — just ignore
        try:
            await cb.message.edit_reply_markup(reply_markup=kb)
        except Exception:
            pass  # Don't send new message — avoid duplicates


def _get_status_text(u: dict) -> str:
    proj = u.get("project_name", "")
    proj_txt = f"📂 {proj}" if proj else "📂 Нет проекта"
    return (
        f"⚖️ <b>MILEAN — Панель управления</b>\n\n"
        f"{proj_txt}\n"
        f"🌐 Web: {'✅ ВКЛ' if u['web_on'] else '❌ ВЫКЛ'} · "
        f"🧠 Think: {'✅ ВКЛ' if u['think_on'] else '❌ ВЫКЛ'}\n"
        f"📝 Инструкция: {'⚖️ MILEAN' if u['active_slot']=='milean' else ('✏️ Своя' if u['instr'] else '⛔ Нет')}\n"
        f"📎 Файлов: {len(u['files'])} · 🧩 Чанков: {len(u['chunks'])}\n"
        f"💬 История: {len(u['hist'])//2} сообщ."
    )


# ─── CALLBACK HANDLERS ───
@router.callback_query(F.data == "toggle_web")
async def cb_web(cb: CallbackQuery):
    u = get_user(cb.from_user.id, cb.bot.id)
    u["web_on"] = not u["web_on"]
    await _update_kb(cb, u, f"🌐 Веб-поиск: {'ВКЛ' if u['web_on'] else 'ВЫКЛ'}")


@router.callback_query(F.data == "toggle_think")
async def cb_think(cb: CallbackQuery):
    u = get_user(cb.from_user.id, cb.bot.id)
    u["think_on"] = not u["think_on"]
    await _update_kb(cb, u, f"🧠 Think: {'ВКЛ' if u['think_on'] else 'ВЫКЛ'}")


@router.callback_query(F.data == "load_milean")
async def cb_milean(cb: CallbackQuery):
    u = get_user(cb.from_user.id, cb.bot.id)
    if u["active_slot"] == "milean":
        # Toggle OFF — clear instruction
        u["instr"] = ""
        u["active_slot"] = "empty"
        await _update_kb(cb, u, "⚖️ MILEAN выключена")
    else:
        # Toggle ON — load MILEAN
        u["instr"] = MILEAN_INSTR
        u["active_slot"] = "milean"
        await _update_kb(cb, u, "⚖️ MILEAN включена")


@router.callback_query(F.data == "clear_instr")
async def cb_clear_instr(cb: CallbackQuery):
    u = get_user(cb.from_user.id, cb.bot.id)
    u["instr"] = ""
    u["active_slot"] = "empty"
    await _update_kb(cb, u, "🗑 Инструкция очищена")


# ─── DOWNLOAD HANDLERS ───
@router.callback_query(F.data.startswith("dl_txt_"))
async def cb_dl_txt(cb: CallbackQuery):
    u = get_user(cb.from_user.id, cb.bot.id)
    content = u.get("last_response", "")
    if not content:
        await cb.answer("❌ Нет ответа для скачивания")
        return
    await cb.answer("📄 Генерация TXT...")
    thinking = u.get("last_thinking", "")
    query = u.get("last_query", "")
    full = ""
    if query:
        full += f"ВОПРОС:\n{query}\n\n{'='*60}\n\n"
    if thinking:
        full += f"РАЗМЫШЛЕНИЕ:\n{thinking}\n\n{'='*60}\n\n"
    full += f"ОТВЕТ:\n{content}"
    buf = BytesIO(full.encode("utf-8"))
    buf.name = "milean_response.txt"
    buf.seek(0)
    await cb.bot.send_document(cb.message.chat.id, types.BufferedInputFile(buf.read(), filename="milean_response.txt"), caption="📄 Ответ MILEAN (TXT)")


@router.callback_query(F.data.startswith("dl_docx_"))
async def cb_dl_docx(cb: CallbackQuery):
    u = get_user(cb.from_user.id, cb.bot.id)
    content = u.get("last_response", "")
    if not content:
        await cb.answer("❌ Нет ответа для скачивания")
        return
    await cb.answer("📝 Генерация DOCX...")
    try:
        from docx import Document as DocxDocument
        from docx.shared import Pt, RGBColor
        doc = DocxDocument()
        # Title
        title = doc.add_heading("MILEAN — Ответ", level=1)
        # Query
        query = u.get("last_query", "")
        if query:
            doc.add_heading("Вопрос", level=2)
            doc.add_paragraph(query)
        # Thinking
        thinking = u.get("last_thinking", "")
        if thinking:
            doc.add_heading("Размышление", level=2)
            p = doc.add_paragraph(thinking)
            for run in p.runs:
                run.font.color.rgb = RGBColor(128, 128, 128)
                run.font.size = Pt(9)
        # Answer
        doc.add_heading("Ответ", level=2)
        for para in content.split("\n"):
            if para.strip():
                doc.add_paragraph(para)
        buf = BytesIO()
        doc.save(buf)
        buf.seek(0)
        await cb.bot.send_document(cb.message.chat.id, types.BufferedInputFile(buf.read(), filename="milean_response.docx"), caption="📝 Ответ MILEAN (DOCX)")
    except Exception as e:
        await cb.message.answer(f"❌ Ошибка генерации DOCX: {e}")


@router.callback_query(F.data.startswith("dl_pdf_"))
async def cb_dl_pdf(cb: CallbackQuery):
    u = get_user(cb.from_user.id, cb.bot.id)
    content = u.get("last_response", "")
    if not content:
        await cb.answer("❌ Нет ответа для скачивания")
        return
    await cb.answer("📕 Генерация PDF...")
    # Generate PDF as TXT fallback (simple approach)
    thinking = u.get("last_thinking", "")
    query = u.get("last_query", "")
    full = ""
    if query:
        full += f"ВОПРОС:\n{query}\n\n{'='*60}\n\n"
    if thinking:
        full += f"РАЗМЫШЛЕНИЕ:\n{thinking}\n\n{'='*60}\n\n"
    full += f"ОТВЕТ:\n{content}"
    try:
        from reportlab.lib.pagesizes import A4
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
        from reportlab.lib.units import mm
        from reportlab.pdfbase import pdfmetrics
        from reportlab.pdfbase.ttfonts import TTFont
        import tempfile
        # Try to use a font that supports Cyrillic
        buf = BytesIO()
        doc_pdf = SimpleDocTemplate(buf, pagesize=A4)
        styles = getSampleStyleSheet()
        story = [Paragraph("MILEAN — Ответ", styles['Title'])]
        for line in full.split("\n"):
            if line.strip():
                story.append(Paragraph(line.replace("&","&amp;").replace("<","&lt;").replace(">","&gt;"), styles['Normal']))
            else:
                story.append(Spacer(1, 3*mm))
        doc_pdf.build(story)
        buf.seek(0)
        await cb.bot.send_document(cb.message.chat.id, types.BufferedInputFile(buf.read(), filename="milean_response.pdf"), caption="📕 Ответ MILEAN (PDF)")
    except ImportError:
        # Fallback: send as TXT with .pdf note
        buf = BytesIO(full.encode("utf-8"))
        await cb.bot.send_document(cb.message.chat.id, types.BufferedInputFile(buf.read(), filename="milean_response.txt"), caption="📄 PDF библиотека не установлена, отправляю TXT")


@router.callback_query(F.data.startswith("dl_chat_"))
async def cb_dl_chat(cb: CallbackQuery):
    u = get_user(cb.from_user.id, cb.bot.id)
    content = u.get("last_response", "")
    if not content:
        await cb.answer("❌ Нет ответа")
        return
    await cb.answer("💬 Отправка в чат...")
    await _send_long(cb.bot, cb.message.chat.id, content)


@router.callback_query(F.data == "show_project")
async def cb_show_project(cb: CallbackQuery):
    u = get_user(cb.from_user.id, cb.bot.id)
    proj_name = u.get("project_name", "—")
    files_info = ""
    if u["files"]:
        files_info = "\n".join(f"  📄 {f['name']} ({f.get('chunks',0)} чанков)" for f in u["files"])
    else:
        files_info = "  нет файлов"
    await cb.answer()
    await cb.message.answer(
        f"📂 <b>Проект: {proj_name}</b>\n\n"
        f"📎 Файлы:\n{files_info}\n\n"
        f"🧩 Чанков: {len(u['chunks'])}\n"
        f"📝 Инструкция: {'✅' if u['instr'] else '❌'}\n"
        f"💬 История: {len(u['hist'])//2} сообщ.",
        parse_mode=ParseMode.HTML
    )


@router.callback_query(F.data == "show_settings")
async def cb_settings(cb: CallbackQuery):
    await cb.answer()
    u = get_user(cb.from_user.id, cb.bot.id)
    await cb.message.answer(
        f"⚙️ Think: {'✅' if u['think_on'] else '❌'} | "
        f"Web: {'✅' if u['web_on'] else '❌'} | "
        f"Файлов: {len(u['files'])} | "
        f"Чанков: {len(u['chunks'])}"
    )


# ─── KEYBOARDS ───
def get_main_keyboard(u: dict) -> InlineKeyboardMarkup:
    web_label = "🌐 Web: ✅" if u["web_on"] else "🌐 Web: ❌"
    think_label = "🧠 Think: ✅" if u["think_on"] else "🧠 Think: ❌"
    instr_label = "⚖️ MILEAN ✅" if u["active_slot"] == "milean" else "⚖️ MILEAN"
    proj_name = u.get("project_name", "")
    proj_label = f"📂 {proj_name}" if proj_name else "📂 Нет проекта"

    rows = [
        [
            InlineKeyboardButton(text=web_label, callback_data="toggle_web"),
            InlineKeyboardButton(text=think_label, callback_data="toggle_think"),
        ],
        [
            InlineKeyboardButton(text=instr_label, callback_data="load_milean"),
            InlineKeyboardButton(text="🗑 Очистить", callback_data="clear_instr"),
        ],
        [
            InlineKeyboardButton(text=proj_label, callback_data="show_project"),
            InlineKeyboardButton(text="⚙️ Статус", callback_data="show_settings"),
        ],
    ]
    return InlineKeyboardMarkup(inline_keyboard=rows)


# ─── HELPERS ───
def _escape(text: str) -> str:
    """Escape HTML special chars"""
    return text.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")


async def _send_long(bot_client: Bot, chat_id: int, text: str, reply_to: int = None):
    """Send long messages by splitting at 4096 chars"""
    MAX_LEN = 4000
    if len(text) <= MAX_LEN:
        try:
            await bot_client.send_message(chat_id, text, parse_mode=ParseMode.HTML, reply_to_message_id=reply_to)
        except:
            # Fallback without HTML if parsing fails
            await bot_client.send_message(chat_id, text, reply_to_message_id=reply_to)
        return

    # Split by paragraphs
    parts = []
    current = ""
    for line in text.split("\n"):
        if len(current) + len(line) + 1 > MAX_LEN:
            if current:
                parts.append(current)
            current = line
        else:
            current += ("\n" if current else "") + line
    if current:
        parts.append(current)

    for i, part in enumerate(parts):
        try:
            await bot_client.send_message(
                chat_id, part,
                parse_mode=ParseMode.HTML,
                reply_to_message_id=reply_to if i == 0 else None
            )
        except:
            await bot_client.send_message(
                chat_id, part,
                reply_to_message_id=reply_to if i == 0 else None
            )
        if i < len(parts) - 1:
            await asyncio.sleep(0.3)


# ─── SETUP BOT COMMANDS ───
async def set_commands(bot_client: Bot):
    commands = [
        BotCommand(command="start", description="🚀 Начать работу"),
        BotCommand(command="help", description="📋 Список команд"),
        BotCommand(command="milean", description="⚖️ Загрузить инструкцию MILEAN"),
        BotCommand(command="instr", description="📝 Показать инструкцию"),
        BotCommand(command="setinstr", description="✏️ Написать свою инструкцию"),
        BotCommand(command="slots", description="📋 Список слотов"),
        BotCommand(command="web", description="🌐 Вкл/выкл веб-поиск"),
        BotCommand(command="think", description="🧠 Вкл/выкл Think"),
        BotCommand(command="files", description="📎 Список файлов"),
        BotCommand(command="clear", description="🔄 Очистить историю"),
        BotCommand(command="panel", description="🎛 Панель управления"),
        BotCommand(command="connect", description="🔗 Подключить проект с web"),
        BotCommand(command="token", description="🔑 Как получить токен"),
        BotCommand(command="codex", description="🤖 Задача для Codex"),
        BotCommand(command="codexstatus", description="📡 Статус очереди Codex"),
        BotCommand(command="settings", description="⚙️ Настройки"),
    ]
    await bot_client.set_my_commands(commands)


async def _prepare_bot(bot_client: Bot):
    # Polling and webhooks are mutually exclusive on Telegram side.
    try:
        await bot_client.delete_webhook(drop_pending_updates=False)
    except Exception as e:
        log.warning(f"Could not delete webhook for bot: {e}")
    await set_commands(bot_client)


def _build_bots() -> list[Bot]:
    return [Bot(token=token) for token in BOT_TOKENS]


# ─── MAIN ───
async def main():
    _validate_env()
    dp.include_router(router)
    bots = _build_bots()
    await asyncio.gather(*(_prepare_bot(bot_client) for bot_client in bots))
    log.info(f"🚀 MILEAN Bot started! Active bots: {len(bots)}")
    await dp.start_polling(*bots)


if __name__ == "__main__":
    try:
        asyncio.run(main())
    except Exception as e:
        _report_exception("bot.main", e)
        raise
