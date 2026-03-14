"""Shared file storage for RAG — uses /tmp on Vercel"""
import json, os, re, io, uuid, time

STORE_DIR = "/tmp/qwen_rag"
META_FILE = os.path.join(STORE_DIR, "meta.json")
CHUNK_SIZE, CHUNK_OVERLAP, TOP_K, MAX_CTX = 600, 100, 4, 4000

def _ensure_dir():
    os.makedirs(STORE_DIR, exist_ok=True)

def _load_meta():
    _ensure_dir()
    if os.path.exists(META_FILE):
        try:
            with open(META_FILE, "r", encoding="utf-8") as f:
                return json.load(f)
        except:
            pass
    return {}

def _save_meta(meta):
    _ensure_dir()
    with open(META_FILE, "w", encoding="utf-8") as f:
        json.dump(meta, f, ensure_ascii=False)

def _content_path(fid):
    return os.path.join(STORE_DIR, f"{fid}.txt")

def get_files():
    meta = _load_meta()
    result = []
    for fid, info in meta.items():
        cp = _content_path(fid)
        chars = 0
        if os.path.exists(cp):
            chars = os.path.getsize(cp)
        result.append({
            "id": fid, "name": info["name"], "size": info["size"],
            "chars": info.get("chars", chars),
            "chunks": info.get("chunks", 0), "enabled": info.get("enabled", True)
        })
    return result

def get_stats():
    meta = _load_meta()
    en = [f for f in meta.values() if f.get("enabled", True)]
    total_chunks = sum(f.get("chunks", 0) for f in en)
    total_chars = sum(f.get("chars", 0) for f in en)
    return {"total_files": len(meta), "enabled_files": len(en), "total_chunks": total_chunks, "total_chars": total_chars}

def chunk_text(text):
    if len(text) <= CHUNK_SIZE:
        return [text] if text.strip() else []
    chunks, start = [], 0
    while start < len(text):
        end = min(start + CHUNK_SIZE, len(text))
        if end < len(text):
            for s in ['\n\n', '. ', '.\n', '! ', '? ']:
                b = text.rfind(s, start + CHUNK_SIZE // 2, end + 50)
                if b > start:
                    end = b + len(s)
                    break
        c = text[start:end].strip()
        if c:
            chunks.append(c)
        start = end - CHUNK_OVERLAP
        if end >= len(text):
            break
    return chunks

def extract_text(fn, raw):
    ext = fn.rsplit(".", 1)[-1].lower() if "." in fn else ""
    if ext == "pdf":
        try:
            import PyPDF2
            pages = [p.extract_text() for p in PyPDF2.PdfReader(io.BytesIO(raw)).pages]
            return "\n\n".join(t.strip() for t in pages if t and t.strip()) or "[PDF empty]"
        except Exception as e:
            return f"[PDF err:{e}]"
    if ext == "docx":
        try:
            import docx
            doc = docx.Document(io.BytesIO(raw))
            p2 = [p.text for p in doc.paragraphs if p.text.strip()]
            for t in doc.tables:
                for r in t.rows:
                    c2 = [c.text.strip() for c in r.cells]
                    if any(c2):
                        p2.append(" | ".join(c2))
            return "\n".join(p2) or "[DOCX empty]"
        except Exception as e:
            return f"[DOCX err:{e}]"
    if ext in ("jpg", "jpeg", "png", "bmp", "tiff", "tif", "webp"):
        return f"[Image:{fn} — OCR not available on Vercel]"
    try:
        return raw.decode("utf-8")
    except:
        pass
    try:
        return raw.decode("latin-1")
    except:
        return f"[bin {len(raw)}B]"

def upload_files(files_list):
    """files_list: list of (filename, raw_bytes)"""
    meta = _load_meta()
    uploaded = []
    for fn, raw in files_list:
        txt = extract_text(fn, raw)
        fid = uuid.uuid4().hex[:8]
        nc = len(chunk_text(txt))
        meta[fid] = {"name": fn, "size": len(raw), "chars": len(txt), "chunks": nc, "enabled": True}
        with open(_content_path(fid), "w", encoding="utf-8") as f:
            f.write(txt)
        uploaded.append({"id": fid, "name": fn, "size": len(raw), "chars": len(txt), "chunks": nc, "enabled": True})
    _save_meta(meta)
    return uploaded

def toggle_file(fid):
    meta = _load_meta()
    if fid in meta:
        meta[fid]["enabled"] = not meta[fid].get("enabled", True)
        _save_meta(meta)
        return {"ok": True, "enabled": meta[fid]["enabled"]}
    return {"error": "not found"}

def toggle_all(enabled):
    meta = _load_meta()
    for f in meta.values():
        f["enabled"] = enabled
    _save_meta(meta)
    return {"ok": True}

def delete_file(fid):
    meta = _load_meta()
    if fid in meta:
        meta.pop(fid)
        cp = _content_path(fid)
        if os.path.exists(cp):
            os.remove(cp)
        _save_meta(meta)
        return {"ok": True}
    return {"error": "not found"}

def search(query):
    t0 = time.time()
    meta = _load_meta()
    # Build chunks from enabled files
    all_chunks = []
    for fid, info in meta.items():
        if not info.get("enabled", True):
            continue
        cp = _content_path(fid)
        if not os.path.exists(cp):
            continue
        with open(cp, "r", encoding="utf-8") as f:
            txt = f.read()
        for c in chunk_text(txt):
            all_chunks.append({"file_id": fid, "file_name": info["name"], "text": c})

    if not all_chunks:
        ms = int((time.time() - t0) * 1000)
        return {"context": "", "chunks_found": 0, "files_used": [], "total_chars": 0, "search_ms": ms}

    # TF-IDF search
    try:
        from sklearn.feature_extraction.text import TfidfVectorizer
        from sklearn.metrics.pairwise import cosine_similarity
        vectorizer = TfidfVectorizer(max_features=15000, ngram_range=(1, 2), sublinear_tf=True)
        matrix = vectorizer.fit_transform([c["text"] for c in all_chunks])
        sc = cosine_similarity(vectorizer.transform([query]), matrix).flatten()
        res, tot, seen = [], 0, set()
        for i in sc.argsort()[::-1]:
            if sc[i] < 0.01 or len(res) >= TOP_K:
                break
            c = all_chunks[i]
            k = c["text"][:80]
            if k in seen:
                continue
            seen.add(k)
            if tot + len(c["text"]) > MAX_CTX:
                break
            res.append({"file_name": c["file_name"], "text": c["text"], "score": float(sc[i])})
            tot += len(c["text"])
    except Exception:
        res = []

    # Fallback: return first chunks if nothing found
    if not res and all_chunks:
        tot = 0
        for c in all_chunks[:3]:
            if tot + len(c["text"]) > MAX_CTX:
                break
            res.append({"file_name": c["file_name"], "text": c["text"], "score": 0.0})
            tot += len(c["text"])

    ctx = "\n\n---\n\n".join(f"[{r['file_name']}]\n{r['text']}" for r in res)
    fu = list(set(r["file_name"] for r in res))
    ms = int((time.time() - t0) * 1000)
    return {"context": ctx, "chunks_found": len(res), "files_used": fu, "total_chars": len(ctx), "search_ms": ms}
